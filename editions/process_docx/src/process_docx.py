from docx import Document

def extract_code_and_text(docx_file='../input/PRINT-EDITION.docx', output_md='../output/PRINT-EDITION.md'):
    # Open the Word document
    doc = Document(docx_file)
    
    # Prepare an empty list to store the processed content
    processed_content = []
    inside_code_block = False  # Track if we are inside a code block

    # Helper function to check if a run is code based on style
    def is_code_style(style_name):
        return "Source Code" in style_name  # Detects 'Source Code' and its variants
    
    # Iterate over paragraphs in the document
    for paragraph in doc.paragraphs:
        paragraph_text = []  # List to accumulate runs within a paragraph
        paragraph_contains_code = False  # To track if the paragraph contains code

        # Iterate over runs in the paragraph
        for run in paragraph.runs:
            style_name = run.style.name
            run_text = run.text.strip()

            # Detect if this run is code by checking the style
            if is_code_style(style_name):
                paragraph_contains_code = True  # This paragraph contains code
                paragraph_text.append(run_text)  # Treat it as part of the code
            else:
                paragraph_text.append(run.text)  # Regular text

        # Join the run texts to form the paragraph's full content
        full_paragraph_text = " ".join(paragraph_text).strip()

        # Rule-based boundary detection for code blocks
        if paragraph_contains_code:
            if not inside_code_block:  # Start a new code block
                processed_content.append('```python')
                inside_code_block = True
            processed_content.append(full_paragraph_text)  # Add code to the block
        elif inside_code_block and not full_paragraph_text:  # Empty line within code block
            processed_content.append('')  # Preserve empty line in code block
        elif inside_code_block:  # Non-code paragraph ends the code block
            processed_content.append('```')  # Close the code block
            inside_code_block = False
            if full_paragraph_text:  # Add the non-code text
                processed_content.append(full_paragraph_text)
        else:
            # Add regular non-code paragraph content
            if full_paragraph_text:
                processed_content.append(full_paragraph_text)

    # Make sure to close any open code block at the end of the document
    if inside_code_block:
        processed_content.append('```')

    # Write the processed content to a Markdown file
    with open(output_md, 'w', encoding='utf-8') as f:
        f.write('\n\n'.join(processed_content))


extract_code_and_text()

# from docx import Document

# def extract_code_and_text(docx_file='../input/PRINT-EDITION.docx', output_md='../output/PRINT-EDITION.md'):
#     # Open the Word document
#     doc = Document(docx_file)
    
#     # Prepare an empty list to store the processed content
#     processed_content = []
    
#     # Iterate over paragraphs in the document
#     for paragraph in doc.paragraphs:
#         paragraph_text = []  # Use a list to accumulate text with proper spacing
#         run_buffer = []      # Buffer to hold code runs
#         inline_code_active = False
#         paragraph_contains_code = False  # To track if the paragraph has inline code

#         # Iterate over runs within the paragraph
#         for run in paragraph.runs:
#             style_name = run.style.name  # Get the style name
#             run_text = run.text.strip()

#             # Handle inline code (any Source Code style variation within text)
#             if "Source Code" in style_name:  # Check if the style name contains 'Source Code'
#                 paragraph_contains_code = True
#                 # If inside a regular paragraph, treat it as inline code
#                 if not run_text:
#                     continue  # Skip empty runs
#                 paragraph_text.append(f"`{run_text}`")  # Treat as inline code
#             else:
#                 # If it's regular text, add it with a space to ensure words don't merge
#                 paragraph_text.append(run.text)

#         # If the paragraph contains only code, treat it as a code block
#         if paragraph_contains_code and not ''.join(paragraph_text).strip():
#             processed_content.append('```python')
#             processed_content.append(''.join(paragraph_text))
#             processed_content.append('```')
#         else:
#             # Add non-code paragraphs (including inline code) with proper spacing
#             if ''.join(paragraph_text).strip():
#                 processed_content.append(' '.join(paragraph_text))

#     # Write the processed content to a Markdown file
#     with open(output_md, 'w', encoding='utf-8') as f:
#         f.write('\n\n'.join(processed_content))
        
# extract_code_and_text('../input/PRINT-EDITION.docx', '../output/PRINT-EDITION.md')